#!/usr/bin/env python3
"""
Optimized script to create enhanced Word documents from markdown content
Faster processing for large files by optimizing tokenization
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def create_enhanced_document(markdown_content, title, output_filename):
    """Create a professionally formatted Word document from markdown content"""
    doc = Document()
    setup_styles(doc)
    parse_markdown_to_docx(doc, markdown_content, title)
    doc.save(output_filename)
    print(f"✅ Created: {output_filename}")

def setup_styles(doc):
    """Configure document styles per QA Enhancement Protocol"""
    styles = doc.styles

    h1 = styles['Heading 1']
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.name = 'Calibri'

    h2 = styles['Heading 2']
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.name = 'Calibri'

    h3 = styles['Heading 3']
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.name = 'Calibri'

    h4 = styles['Heading 4']
    h4.font.size = Pt(11)
    h4.font.bold = True
    h4.font.name = 'Calibri'

    normal = styles['Normal']
    normal.font.size = Pt(11)
    normal.font.name = 'Calibri'
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

def parse_markdown_to_docx(doc, content, title):
    """Parse markdown and add to document with proper formatting"""
    try:
        doc.add_heading(title, level=1)
        lines = content.split('\n')
        i = 0

        while i < len(lines):
            try:
                line = lines[i].rstrip()

                if not line and i == 0:
                    i += 1
                    continue

                if line.startswith('# ') and title in line:
                    i += 1
                    continue

                # Horizontal rules
                if line.strip() in ['---', '***', '___']:
                    add_horizontal_rule(doc)
                    i += 1
                    continue

                # Headers
                if line.startswith('#### '):
                    doc.add_heading(line[5:], level=4)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('# '):
                    doc.add_heading(line[2:], level=1)

                # Tables
                elif line.startswith('|') and i + 1 < len(lines) and lines[i+1].startswith('|'):
                    i = add_table(doc, lines, i)
                    continue

                # Lists
                elif line.startswith('- ') or line.startswith('  -') or line.startswith('    -') or re.match(r'^(\s*)\d+\.\s', line):
                    i = add_list(doc, lines, i)
                    continue

                # Code blocks
                elif line.startswith('```'):
                    i = add_code_block(doc, lines, i)
                    continue

                # Regular paragraphs
                elif line:
                    p = doc.add_paragraph()
                    # Check if line has markdown formatting before processing
                    if '*' in line or '`' in line:
                        add_formatted_text(p, line)
                    else:
                        p.add_run(line)

                i += 1

            except Exception as e:
                print(f"⚠️  Warning: Error processing line {i}: {str(e)}")
                i += 1
                continue

    except Exception as e:
        print(f"❌ Error during document parsing: {str(e)}")
        raise

def add_horizontal_rule(doc):
    """Add a horizontal rule to the document"""
    p = doc.add_paragraph()
    p_format = p.paragraph_format
    p_format.space_before = Pt(6)
    p_format.space_after = Pt(6)

    p_element = p._element
    p_pr = p_element.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')

    p_bdr.append(bottom)
    p_pr.append(p_bdr)

def add_table(doc, lines, start_idx):
    """Parse and add a markdown table"""
    try:
        table_lines = []
        i = start_idx

        while i < len(lines) and lines[i].startswith('|'):
            table_lines.append(lines[i])
            i += 1

        if len(table_lines) < 2:
            return i

        headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
        rows = []

        for line in table_lines[2:]:
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                rows.append(cells)

        table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Light Grid Accent 1'

        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = ''
            if '*' in header or '`' in header:
                add_formatted_text(cell.paragraphs[0], header)
            else:
                cell.paragraphs[0].add_run(header)
            for run in cell.paragraphs[0].runs:
                run.font.bold = True

        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.rows[row_idx + 1].cells[col_idx]
                cell.text = ''
                if '*' in cell_data or '`' in cell_data:
                    add_formatted_text(cell.paragraphs[0], cell_data)
                else:
                    cell.paragraphs[0].add_run(cell_data)

        doc.add_paragraph()
        return i

    except Exception as e:
        print(f"⚠️  Warning: Error parsing table: {str(e)}")
        return start_idx + 1

def add_list(doc, lines, start_idx):
    """Parse and add a list with nesting support"""
    try:
        i = start_idx

        while i < len(lines):
            line = lines[i].rstrip()

            if not line:
                break

            indent_level = 0
            stripped = line.lstrip()
            leading_spaces = len(line) - len(stripped)

            if leading_spaces >= 2:
                indent_level = leading_spaces // 2

            if stripped.startswith('- '):
                text = stripped[2:]
                p = doc.add_paragraph(style='List Bullet')
                if '*' in text or '`' in text:
                    add_formatted_text(p, text)
                else:
                    p.add_run(text)
                if indent_level > 0:
                    p.paragraph_format.left_indent = Inches(0.25 * indent_level)

            elif re.match(r'^\d+\.\s', stripped):
                match = re.match(r'^\d+\.\s(.+)', stripped)
                if match:
                    text = match.group(1)
                    p = doc.add_paragraph(style='List Number')
                    if '*' in text or '`' in text:
                        add_formatted_text(p, text)
                    else:
                        p.add_run(text)
                    if indent_level > 0:
                        p.paragraph_format.left_indent = Inches(0.25 * indent_level)
            else:
                break

            i += 1

        return i

    except Exception as e:
        print(f"⚠️  Warning: Error parsing list: {str(e)}")
        return start_idx + 1

def add_code_block(doc, lines, start_idx):
    """Add a code block with monospace font and shading"""
    try:
        i = start_idx + 1
        code_lines = []

        while i < len(lines) and not lines[i].startswith('```'):
            code_lines.append(lines[i])
            i += 1

        for code_line in code_lines:
            p = doc.add_paragraph()
            run = p.add_run(code_line)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)

            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F0F0F0')
            p._element.get_or_add_pPr().append(shading_elm)

            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0

        doc.add_paragraph()
        return i + 1 if i < len(lines) else i

    except Exception as e:
        print(f"⚠️  Warning: Error parsing code block: {str(e)}")
        return start_idx + 1

def add_formatted_text(paragraph, text):
    """Add text with markdown formatting support"""
    try:
        tokens = tokenize_markdown_fast(text)

        for token in tokens:
            token_type = token['type']
            content = token['content']

            if token_type == 'bold':
                run = paragraph.add_run(content)
                run.font.bold = True
            elif token_type == 'italic':
                run = paragraph.add_run(content)
                run.font.italic = True
            elif token_type == 'bold_italic':
                run = paragraph.add_run(content)
                run.font.bold = True
                run.font.italic = True
            elif token_type == 'code':
                run = paragraph.add_run(content)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'F5F5F5')
                run._element.get_or_add_rPr().append(shading_elm)
            else:
                paragraph.add_run(content)

    except Exception as e:
        print(f"⚠️  Warning: Error parsing formatted text: {str(e)}")
        paragraph.add_run(text)

def tokenize_markdown_fast(text):
    """Fast tokenization without slow operations"""
    tokens = []
    i = 0

    while i < len(text):
        # Check for bold+italic
        if text[i:i+3] == '***':
            end = text.find('***', i + 3)
            if end != -1:
                tokens.append({'type': 'bold_italic', 'content': text[i+3:end]})
                i = end + 3
                continue

        # Check for bold
        if text[i:i+2] == '**':
            end = text.find('**', i + 2)
            if end != -1:
                tokens.append({'type': 'bold', 'content': text[i+2:end]})
                i = end + 2
                continue

        # Check for italic
        if text[i] == '*' and (i == 0 or text[i-1] != '*'):
            end = text.find('*', i + 1)
            if end != -1 and (end + 1 >= len(text) or text[end+1] != '*'):
                tokens.append({'type': 'italic', 'content': text[i+1:end]})
                i = end + 1
                continue

        # Check for inline code
        if text[i] == '`':
            end = text.find('`', i + 1)
            if end != -1:
                tokens.append({'type': 'code', 'content': text[i+1:end]})
                i = end + 1
                continue

        # Find next special character
        next_pos = len(text)
        for j in range(i + 1, len(text)):
            if text[j] in ['*', '`']:
                next_pos = j
                break

        if next_pos == len(text):
            tokens.append({'type': 'text', 'content': text[i:]})
            break
        else:
            tokens.append({'type': 'text', 'content': text[i:next_pos]})
            i = next_pos

    return tokens


if __name__ == '__main__':
    import sys

    print("Enhanced Word Document Creator (Optimized)")
    print("=" * 50)

    if len(sys.argv) >= 3:
        input_file = sys.argv[1]
        output_file = sys.argv[2]
    else:
        print("Usage: python create_enhanced_docx_optimized.py <input_markdown> <output_docx>")
        sys.exit(1)

    try:
        with open(input_file, 'r', encoding='utf-8') as f:
            markdown_content = f.read()

        title = "Document"
        for line in markdown_content.split('\n'):
            if line.startswith('# '):
                title = line[2:].strip()
                break

        create_enhanced_document(markdown_content, title, output_file)
        print(f"✅ Successfully converted: {input_file} -> {output_file}")

    except FileNotFoundError:
        print(f"❌ Error: Input file not found: {input_file}")
        sys.exit(1)
    except Exception as e:
        print(f"❌ Error: {str(e)}")
        sys.exit(1)
